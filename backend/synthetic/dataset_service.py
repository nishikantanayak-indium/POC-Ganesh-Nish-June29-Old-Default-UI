"""
Synthetic Dataset Management Service — the orchestrator (core #4).

Responsibilities:
  * run the generate → validate → quality → stage pipeline (with a bounded
    regeneration loop for records that fail validation),
  * generate relationship examples + assemble composite documents,
  * snapshot artifacts to the object store and record version stats,
  * maintain lineage,
  * promote a version staging → main,
  * publish a main version into an Analysis workspace (feeding the existing
    knowledge-graph pipeline unchanged).

``run_generation`` is a plain synchronous method driven by a ``progress_cb``
so the API layer can stream progress over SSE (mirroring the Analysis
pipeline route). Keeping orchestration here (not in the route) keeps it
testable and easy to evolve.
"""
from __future__ import annotations

import concurrent.futures
import io
import json
import logging
import random
import uuid
import zipfile
from collections import Counter
from dataclasses import replace
from datetime import datetime, timezone
from typing import Callable, Dict, List, Optional

from core.models import (
    AtomicElement,
    DocumentType,
    ElementType,
    ParsedDocument,
    Relationship,
)

from . import db, taxonomy
from .generation_service import (
    STRUCTURE_HINT_FULL_TEXT_MAX_PAGES,
    DealContext,
    StructureHint,
    SyntheticDataGenerationService,
    _new_id,
)
from .models import (
    DatasetStatus,
    MatrixCell,
    QualityReport,
    RecordStatus,
    SyntheticDocument,
    SyntheticRecord,
    SyntheticRelationship,
    ValidationReport,
)
from .quality_service import SyntheticDataQualityAssessmentService
from .storage import get_artifact_store
from .validation_service import SyntheticDataValidationService

logger = logging.getLogger(__name__)

ProgressCB = Callable[[dict], None]


def _noop(_: dict) -> None:  # pragma: no cover
    pass


def _with_heartbeat(fn: Callable[[], object], emit: Callable[[str], None],
                     running_message: str, interval: float = 2.0) -> object:
    """Run a blocking call (an LLM request) on a worker thread while the
    caller keeps emitting periodic progress messages — otherwise a 10-60s
    OpenAI call reads as total silence to anyone watching the SSE stream."""
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
        fut = ex.submit(fn)
        elapsed = 0.0
        while True:
            try:
                return fut.result(timeout=interval)
            except concurrent.futures.TimeoutError:
                elapsed += interval
                emit(f"{running_message} ({int(elapsed)}s elapsed)")


class SyntheticDatasetManagementService:
    def __init__(
        self,
        generation: SyntheticDataGenerationService,
        validation: SyntheticDataValidationService,
        quality: SyntheticDataQualityAssessmentService,
    ) -> None:
        self.gen = generation
        self.val = validation
        self.qual = quality
        self.store = get_artifact_store()

    # ==================================================================
    # Generation pipeline
    # ==================================================================

    def run_generation(
        self,
        project_id: str,
        selections: List[dict],           # [{"cell": "Clause|Legal", "count": 5}, ...]
        knobs: Optional[dict] = None,
        progress_cb: ProgressCB = _noop,
    ) -> dict:
        knobs = knobs or {}
        project = db.get_project(project_id)
        if project is None:
            raise ValueError(f"project {project_id} not found")

        industries = knobs.get("industries")
        languages = knobs.get("languages")
        doc_types = [DocumentType(d) for d in knobs.get("doc_types", [])] or None
        do_relationships = knobs.get("generate_relationships", True)
        do_documents = knobs.get("assemble_documents", True)
        # Default: fold every staged record into ONE composite document, even
        # when the batch spans several doc_types (doc_type stays a content
        # diversity dimension, not a file-splitting key). Users who explicitly
        # want one file per doc_type can opt in.
        split_by_doc_type = knobs.get("split_by_doc_type", False)
        brief = knobs.get("brief") or ""
        mirror_document_id = knobs.get("mirror_document_id")
        max_regen = knobs.get("max_regen")
        from config.settings import settings
        if max_regen is None:
            max_regen = settings.synthetic_max_regen

        seed_examples: Dict[str, List[str]] = (project.seed_summary or {}).get("examples", {})

        # ── Mirror mode: reproduce a specific seed document's shape ──────────
        mirror_doc: Optional[dict] = None
        if mirror_document_id:
            for d in (project.seed_summary or {}).get("documents", []):
                if d.get("id") == mirror_document_id:
                    mirror_doc = d
                    break
            if mirror_doc:
                # Prefer the mirrored document's own examples for few-shot conditioning.
                if mirror_doc.get("examples"):
                    seed_examples = {**seed_examples, **mirror_doc["examples"]}
                # If the caller didn't specify selections, derive them from the doc's composition.
                if not selections and mirror_doc.get("cells"):
                    selections = [{"cell": k, "count": int(v)} for k, v in mirror_doc["cells"].items()]
                progress_cb({"stage": "start",
                             "message": f"Mirroring '{mirror_doc.get('name')}' — {len(selections)} cells"})

        dataset = db.get_or_create_default_dataset(project_id)
        version = db.create_version(project_id, dataset.id, note=knobs.get("note", ""))
        progress_cb({"stage": "start", "message": f"Version v{version.version_no} created",
                     "version_id": version.id})

        labels = project.label_set
        all_records: List[SyntheticRecord] = []
        all_reports = []  # ValidationReport
        total_target = sum(int(s["count"]) for s in selections)
        done = 0

        # ── generate + validate (+ bounded regeneration) per selection ────
        # A selection is either cell-based {"cell": "Type|Label"} (Balance mode,
        # fixed label) or type-only {"element_type": "Clause"} (Describe mode,
        # the model assigns the best label per record).
        for sel in selections:
            count = int(sel["count"])
            if count <= 0:
                continue
            if sel.get("cell"):
                cell = MatrixCell.from_key(sel["cell"])
                et, fixed_label, seeds, tag = cell.element_type, cell.label, seed_examples.get(cell.key), cell.key
            else:
                et = ElementType(sel["element_type"])
                fixed_label, seeds, tag = None, None, f"{et.value}|auto"

            progress_cb({"stage": "generate", "message": f"Generating {count}× {tag}",
                         "current": done, "total": total_target, "cell": tag})

            passing: List[SyntheticRecord] = []
            attempts = 0
            need = count
            while need > 0 and attempts <= max_regen:
                batch = self.gen.generate_records(
                    project_id, et, need, label=fixed_label, allowed_labels=labels, seeds=seeds,
                    industries=industries, languages=languages, doc_types=doc_types,
                    version_id=version.id, brief=brief,
                )
                reports = self.val.validate_records(batch, allowed_labels=labels)
                rep_by_id = {r.record_id: r for r in reports}
                for rec in batch:
                    rep = rep_by_id[rec.id]
                    if rep.passed:
                        passing.append(rec)
                        all_reports.append(rep)
                    else:
                        rec.status = RecordStatus.REJECTED
                        all_records.append(rec)      # keep rejected for transparency
                        all_reports.append(rep)
                need = count - len(passing)
                attempts += 1
                if need > 0 and attempts <= max_regen:
                    progress_cb({"stage": "validate",
                                 "message": f"{tag}: {need} failed validation — regenerating (attempt {attempts})",
                                 "current": done, "total": total_target, "cell": tag})

            all_records.extend(passing)
            done += count
            progress_cb({"stage": "validate", "message": f"{tag}: {len(passing)} valid",
                         "current": done, "total": total_target, "cell": tag})

        valid_records = [r for r in all_records if r.status != RecordStatus.REJECTED]

        # ── quality assessment (dedup + realism) ─────────────────────────
        progress_cb({"stage": "quality", "message": f"Assessing quality of {len(valid_records)} records"})
        quality_reports = self.qual.assess_records(valid_records)
        qr_by_id = {q.record_id: q for q in quality_reports}
        staged: List[SyntheticRecord] = []
        for rec in valid_records:
            q = qr_by_id.get(rec.id)
            if q and q.passed:
                rec.status = RecordStatus.STAGED
                staged.append(rec)
            elif q and q.is_duplicate:
                rec.status = RecordStatus.DUPLICATE
            else:
                rec.status = RecordStatus.REJECTED
        progress_cb({"stage": "quality",
                     "message": f"{len(staged)} staged · "
                                f"{sum(1 for q in quality_reports if q.is_duplicate)} duplicates · "
                                f"{sum(1 for q in quality_reports if not q.passed and not q.is_duplicate)} low-realism"})

        # ── relationships ────────────────────────────────────────────────
        relationships: List[SyntheticRelationship] = []
        if do_relationships and staged:
            progress_cb({"stage": "relate", "message": "Generating relationship / mapping examples"})
            rels = self.gen.generate_relationships(project_id, staged, version_id=version.id)
            id_to_type = {r.id: r.element_type for r in staged}
            reasons = self.val.validate_relationships(rels, id_to_type)
            for rel in rels:
                if not reasons.get(rel.id):
                    rel.status = RecordStatus.STAGED
                    relationships.append(rel)
            progress_cb({"stage": "relate",
                         "message": f"{len(relationships)}/{len(rels)} relationships passed coverage consistency"})

        # ── composite documents ─────────────────────────────────────────
        documents: List[SyntheticDocument] = []
        if do_documents and staged:
            if mirror_doc:
                # Single document mirroring the source's type + section layout.
                progress_cb({"stage": "assemble",
                             "message": f"Assembling 1 document mirroring '{mirror_doc.get('name')}'"})
                try:
                    dtype = DocumentType(mirror_doc.get("type", staged[0].doc_type.value))
                except ValueError:
                    dtype = staged[0].doc_type
                doc, markdown = self.gen.assemble_document(
                    project_id, staged, dtype, version_id=version.id,
                    structure=mirror_doc.get("sections"), brief=brief,
                )
                key = self._doc_key(project_id, dataset.id, version.version_no, doc.id)
                doc.artifact_uri = self.store.put_text(key, markdown, "text/markdown")
                doc.provenance["artifact_key"] = key
                documents.append(doc)
            elif split_by_doc_type:
                progress_cb({"stage": "assemble", "message": "Assembling one document per document type"})
                by_doc: Dict[DocumentType, List[SyntheticRecord]] = {}
                for r in staged:
                    by_doc.setdefault(r.doc_type, []).append(r)
                for dtype, recs in by_doc.items():
                    doc, markdown = self.gen.assemble_document(
                        project_id, recs, dtype, version_id=version.id, brief=brief,
                    )
                    key = self._doc_key(project_id, dataset.id, version.version_no, doc.id)
                    doc.artifact_uri = self.store.put_text(key, markdown, "text/markdown")
                    doc.provenance["artifact_key"] = key
                    documents.append(doc)
            else:
                progress_cb({"stage": "assemble", "message": "Assembling 1 composite document"})
                # doc_type stays a per-record diversity dimension; the document's
                # own doc_type is just the most common one in this batch.
                dtype = Counter(r.doc_type for r in staged).most_common(1)[0][0]
                doc, markdown = self.gen.assemble_document(
                    project_id, staged, dtype, version_id=version.id, brief=brief,
                )
                key = self._doc_key(project_id, dataset.id, version.version_no, doc.id)
                doc.artifact_uri = self.store.put_text(key, markdown, "text/markdown")
                doc.provenance["artifact_key"] = key
                documents.append(doc)

        # ── persist everything ───────────────────────────────────────────
        progress_cb({"stage": "persist", "message": "Persisting records, reports, artifacts"})
        db.insert_records(all_records)
        db.upsert_validation_reports(all_reports)
        db.upsert_quality_reports(quality_reports)
        db.insert_relationships(relationships)
        for d in documents:
            db.insert_document(d)

        # ── snapshot + stats + lineage ───────────────────────────────────
        distribution = self.qual.compute_distribution(staged, relationships, project.min_threshold)
        artifact_uri = self._snapshot(project_id, dataset.id, version.version_no,
                                      staged, relationships)
        stats = {
            "requested": total_target,
            "generated": len(all_records),
            "staged": len(staged),
            "rejected": sum(1 for r in all_records if r.status == RecordStatus.REJECTED),
            "duplicates": sum(1 for r in all_records if r.status == RecordStatus.DUPLICATE),
            "relationships": len(relationships),
            "documents": len(documents),
            "distribution": distribution,
        }
        db.update_version_stats(version.id, stats, artifact_uri)
        self._lineage_for_generation(project_id, version.id, selections, staged)
        db.touch_project(project_id)

        summary = {"version_id": version.id, "version_no": version.version_no, **stats}
        progress_cb({"stage": "complete", "message": "Generation complete", "summary": summary})
        return summary

    # ==================================================================
    # Document-level generation (document-first pivot)
    # ==================================================================

    def run_document_generation(
        self,
        project_id: str,
        doc_targets: List[dict],   # [{"doc_type": "RFP", "count": 5, "brief": "..."}, ...]
        knobs: Optional[dict] = None,
        progress_cb: ProgressCB = _noop,
    ) -> dict:
        """
        Generate whole documents directly — no atomic elements, no LLM
        validation/quality gating (SME review is the sole quality gate for
        this flow). Deliberately does not touch ``run_generation`` above,
        which remains the parked element-level path.
        """
        knobs = knobs or {}
        project = db.get_project(project_id)
        if project is None:
            raise ValueError(f"project {project_id} not found")

        industries = knobs.get("industries")
        languages = knobs.get("languages")
        note = knobs.get("note", "")
        mode = knobs.get("mode", "independent")
        seed_docs = (project.seed_summary or {}).get("documents", [])

        dataset = db.get_or_create_default_dataset(project_id)
        version = db.create_version(project_id, dataset.id, note=note)
        progress_cb({"stage": "start", "message": f"Version v{version.version_no} created",
                     "version_id": version.id})

        documents: List[SyntheticDocument] = []
        by_doc_type: Dict[str, int] = {}

        def _seeds_for(dtype: DocumentType) -> List[str]:
            # Reuse any matching seed document's captured examples for tone
            # conditioning, same convention as mirror-mode in run_generation.
            out: List[str] = []
            for d in seed_docs:
                if d.get("type") == dtype.value and d.get("examples"):
                    for vals in d["examples"].values():
                        out.extend(vals)
            return out

        def _structure_hint_for(dtype: DocumentType) -> Optional[StructureHint]:
            # Ground generation in a real uploaded document's structure — full
            # text for short documents (fits the prompt comfortably), just the
            # section-heading order for long ones (already captured at upload).
            match = next((d for d in seed_docs if d.get("type") == dtype.value), None)
            if match is None:
                return None
            total_pages = int(match.get("total_pages") or 0)
            if total_pages and total_pages <= STRUCTURE_HINT_FULL_TEXT_MAX_PAGES and match.get("text_key"):
                try:
                    full_text = self.store.get_bytes(match["text_key"]).decode("utf-8")
                    if full_text.strip():
                        return StructureHint(source_name=match.get("name", dtype.value), full_text=full_text)
                except Exception as exc:
                    logger.warning("structure hint: failed to load full text for %s: %s", match.get("id"), exc)
            headings = [s.get("heading") for s in match.get("sections", []) if s.get("heading")]
            if headings:
                return StructureHint(source_name=match.get("name", dtype.value), section_headings=headings)
            return None

        if mode == "linked":
            deal_count = max(1, int(knobs.get("deal_count", 1)))
            total_target = deal_count * 3  # RFP + Contract + RiskSheet per deal
            done = 0
            for deal_no in range(1, deal_count + 1):
                def _deal_progress(msg: str, current: Optional[int] = None, stage: str = "generate") -> None:
                    progress_cb({
                        "stage": stage, "message": f"Deal {deal_no}/{deal_count}: {msg}",
                        "current": done + (current or 0), "total": total_target, "cell": "deal",
                    })

                deal_docs = self._generate_linked_deal(
                    project_id, version.id, industries, languages, note,
                    seeds_for=_seeds_for, structure_hint_for=_structure_hint_for,
                    progress_cb=_deal_progress,
                )
                for doc in deal_docs:
                    _deal_progress(f"writing {doc.doc_type.value} to storage…", stage="persist")
                    key = self._doc_key(project_id, dataset.id, version.version_no, doc.id)
                    # markdown is stashed in provenance by _generate_linked_deal, pop it back out
                    markdown = doc.provenance.pop("_markdown")
                    doc.artifact_uri = self.store.put_text(key, markdown, "text/markdown")
                    doc.provenance["artifact_key"] = key
                    documents.append(doc)
                    by_doc_type[doc.doc_type.value] = by_doc_type.get(doc.doc_type.value, 0) + 1
                    done += 1
        else:
            total_target = sum(int(t["count"]) for t in doc_targets)
            done = 0

            for target in doc_targets:
                dtype = DocumentType(target["doc_type"])
                count = int(target.get("count", 0))
                if count <= 0:
                    continue
                brief = target.get("brief") or ""
                seeds = _seeds_for(dtype)
                hint = _structure_hint_for(dtype)

                for i in range(count):
                    base_evt = {"current": done, "total": total_target, "cell": dtype.value}

                    def _emit(message: str, stage: str = "generate") -> None:
                        progress_cb({"stage": stage, "message": message, **base_evt})

                    _emit(
                        f"Preparing {dtype.value} {i + 1}/{count} — brief: "
                        f"{'yes' if brief.strip() else 'no'}, structure template: "
                        f"{hint.source_name if hint else 'none'}, seed tone examples: {len(seeds)}"
                    )
                    doc, markdown, _facts = _with_heartbeat(
                        lambda: self.gen.generate_document(
                            project_id, dtype, version_id=version.id,
                            seeds=seeds or None, industries=industries, languages=languages, brief=brief,
                            note=note, structure_hint=hint,
                        ),
                        _emit, f"Drafting {dtype.value} {i + 1}/{count} via LLM",
                    )
                    word_count = len(markdown.split())
                    _emit(f"{dtype.value} drafted — {len(doc.sections)} sections, ~{word_count} words")

                    dims_preview = ["realism"] + (["structural fidelity"] if hint else []) + \
                        (["instruction adherence"] if (brief.strip() or note.strip()) else [])
                    _emit(f"Validating {dtype.value} — checking {', '.join(dims_preview)}…", stage="validate")
                    validation = _with_heartbeat(
                        lambda: self.gen.validate_document(
                            dtype, markdown, brief=brief, note=note, structure_hint=hint,
                        ),
                        lambda m: _emit(m, stage="validate"), f"Validating {dtype.value}",
                    )
                    doc.provenance["validation"] = validation
                    if validation.get("error"):
                        _emit(f"{dtype.value}: validation failed ({validation['error']})", stage="validate")
                    else:
                        pct = round((validation.get("overall_score") or 0) * 100)
                        _emit(f"{dtype.value} validated — overall score {pct}%", stage="validate")

                    _emit(f"Writing {dtype.value} to storage…", stage="persist")
                    key = self._doc_key(project_id, dataset.id, version.version_no, doc.id)
                    doc.artifact_uri = self.store.put_text(key, markdown, "text/markdown")
                    doc.provenance["artifact_key"] = key
                    documents.append(doc)
                    by_doc_type[dtype.value] = by_doc_type.get(dtype.value, 0) + 1
                    done += 1
                    progress_cb({"stage": "generate", "message": f"{dtype.value}: {done}/{total_target} generated",
                                 "current": done, "total": total_target, "cell": dtype.value})

        # ── persist ────────────────────────────────────────────────────
        progress_cb({"stage": "persist", "message": "Persisting documents"})
        for d in documents:
            db.insert_document(d)

        distribution = {
            dtype.value: {
                "generated": by_doc_type.get(dtype.value, 0),
                "threshold": project.min_threshold,
            }
            for dtype in taxonomy.DEFAULT_DOC_TYPES
        }
        stats = {
            "requested": total_target,
            "generated": len(documents),
            "staged": len(documents),
            "documents": len(documents),
            "distribution": distribution,
        }
        db.update_version_stats(version.id, stats)
        db.add_lineage_edges(project_id, [
            (f"project:{project_id}", f"version:{version.id}", "generated"),
        ])
        db.touch_project(project_id)

        summary = {"version_id": version.id, "version_no": version.version_no, **stats}
        progress_cb({"stage": "complete", "message": "Generation complete", "summary": summary})
        return summary

    def _generate_linked_deal(
        self,
        project_id: str,
        version_id: str,
        industries: Optional[List[str]],
        languages: Optional[List[str]],
        note: str,
        seeds_for: Callable[[DocumentType], List[str]],
        progress_cb: Callable[..., None],
        structure_hint_for: Optional[Callable[[DocumentType], Optional[StructureHint]]] = None,
        covered_ratio: float = 0.75,
    ) -> List[SyntheticDocument]:
        """
        Generate one RFP -> Contract -> RiskSheet deal where each document
        genuinely references concrete facts from the one before it, so the
        real (unmodified) cross-document extractor finds real relationships
        on import — see plan doc for rationale. Deliberately holds back a
        minority of facts at each step for realistic coverage/mitigation gaps.

        Returns the 3 documents with their markdown temporarily stashed at
        ``provenance["_markdown"]`` (popped by the caller before persisting)
        since this method doesn't have access to the artifact store key
        scheme (that's assembled per-version by the caller).
        """
        deal_id = f"DEAL_{uuid.uuid4().hex[:8].upper()}"

        def _split(facts: List[Dict[str, str]]) -> tuple[List[str], List[str]]:
            texts = [f["text"] for f in facts]
            random.shuffle(texts)
            n_covered = max(1, round(len(texts) * covered_ratio)) if texts else 0
            return texts[:n_covered], texts[n_covered:]

        def _draft(step_no: int, dtype: DocumentType, hint: Optional[StructureHint], **gen_kwargs):
            deal_ctx = gen_kwargs.get("deal_context")
            progress_cb(
                f"preparing {dtype.value} — structure template: {hint.source_name if hint else 'none'}, "
                f"deal context: {'yes' if deal_ctx else 'no (source document)'}", step_no,
            )
            doc, md, facts = _with_heartbeat(
                lambda: self.gen.generate_document(
                    project_id, dtype, version_id=version_id,
                    industries=industries, languages=languages, note=note,
                    structure_hint=hint, **gen_kwargs,
                ),
                lambda m: progress_cb(m, step_no), f"Drafting {dtype.value} via LLM",
            )
            word_count = len(md.split())
            progress_cb(f"{dtype.value} drafted — {len(doc.sections)} sections, ~{word_count} words", step_no)

            dims_preview = ["realism"] + (["structural fidelity"] if hint else []) + \
                (["deal consistency"] if deal_ctx else [])
            progress_cb(f"Validating {dtype.value} — checking {', '.join(dims_preview)}…", step_no, "validate")
            validation = _with_heartbeat(
                lambda: self.gen.validate_document(
                    dtype, md, note=note, structure_hint=hint, deal_context=deal_ctx,
                ),
                lambda m: progress_cb(m, step_no, "validate"), f"Validating {dtype.value}",
            )
            doc.provenance["validation"] = validation
            if validation.get("error"):
                progress_cb(f"{dtype.value}: validation failed ({validation['error']})", step_no, "validate")
            else:
                pct = round((validation.get("overall_score") or 0) * 100)
                progress_cb(f"{dtype.value} validated — overall score {pct}%", step_no, "validate")
            return doc, md, facts

        # 1. RFP — the source of the deal's concrete requirements.
        rfp_doc, rfp_md, rfp_facts = _draft(
            1, DocumentType.RFP, structure_hint_for(DocumentType.RFP) if structure_hint_for else None,
            seeds=seeds_for(DocumentType.RFP) or None, emit_key_facts=True,
        )
        rfp_doc.provenance.update({"deal_id": deal_id, "deal_role": "source"})
        rfp_doc.provenance["_markdown"] = rfp_md
        rfp_covered, rfp_held_back = _split(rfp_facts)

        # 2. Contract — responds to the RFP, covering most (not all) of its facts.
        contract_doc, contract_md, contract_facts = _draft(
            2, DocumentType.CONTRACT, structure_hint_for(DocumentType.CONTRACT) if structure_hint_for else None,
            seeds=seeds_for(DocumentType.CONTRACT) or None,
            deal_context=DealContext("RFP", covered_facts=rfp_covered, held_back_facts=rfp_held_back),
            emit_key_facts=True,
        )
        contract_doc.provenance.update({
            "deal_id": deal_id, "deal_role": "downstream",
            "deal_uncovered_requirements": rfp_held_back,
        })
        contract_doc.provenance["_markdown"] = contract_md
        contract_covered, contract_held_back = _split(contract_facts)

        # 3. RiskSheet — analyzes the Contract's clauses, leaving some risks unmitigated.
        risk_doc, risk_md, _risk_facts = _draft(
            3, DocumentType.RISK_SHEET, structure_hint_for(DocumentType.RISK_SHEET) if structure_hint_for else None,
            seeds=seeds_for(DocumentType.RISK_SHEET) or None,
            deal_context=DealContext(
                "Contract", covered_facts=contract_covered, held_back_facts=contract_held_back,
            ),
        )
        risk_doc.provenance.update({
            "deal_id": deal_id, "deal_role": "downstream",
            "deal_unmitigated_facts": contract_held_back,
        })
        risk_doc.provenance["_markdown"] = risk_md

        return [rfp_doc, contract_doc, risk_doc]

    def doc_type_overview(self, project_id: str) -> dict:
        """Per-document-type gap analysis: how many exist (seed + generated)
        vs. the project's threshold. The doc-level counterpart to the
        ElementType×Label matrix overview used by the parked element flow."""
        project = db.get_project(project_id)
        if project is None:
            raise ValueError(f"project {project_id} not found")

        seed_docs = (project.seed_summary or {}).get("documents", [])
        seed_counts: Dict[str, int] = {}
        for d in seed_docs:
            t = d.get("type")
            if t:
                seed_counts[t] = seed_counts.get(t, 0) + 1

        generated = db.list_project_documents(
            project_id,
            statuses=[RecordStatus.STAGED, RecordStatus.SME_APPROVED, RecordStatus.PUBLISHED],
        )
        generated_counts: Dict[str, int] = {}
        for d in generated:
            generated_counts[d.doc_type.value] = generated_counts.get(d.doc_type.value, 0) + 1

        types = []
        for dtype in taxonomy.DEFAULT_DOC_TYPES:
            seed_count = seed_counts.get(dtype.value, 0)
            gen_count = generated_counts.get(dtype.value, 0)
            total = seed_count + gen_count
            types.append({
                "doc_type": dtype.value,
                "seed_count": seed_count,
                "generated_count": gen_count,
                "total": total,
                "threshold": project.min_threshold,
                "deficit": max(0, project.min_threshold - total),
            })
        return {"project_id": project_id, "min_threshold": project.min_threshold, "doc_types": types}

    # ==================================================================
    # Promotion + publication
    # ==================================================================

    def promote(self, version_id: str) -> dict:
        version = db.get_version(version_id)
        if version is None:
            raise ValueError(f"version {version_id} not found")
        db.set_version_status(version_id, DatasetStatus.MAIN)
        db.add_lineage_edges(version.project_id, [
            (f"version:{version_id}", f"main:{version_id}", "promoted"),
        ])
        return {"version_id": version_id, "status": DatasetStatus.MAIN.value}

    def delete_version(self, version_id: str) -> dict:
        """Delete a version and its artifacts. Records/rels/docs/reports cascade
        in Postgres; the object-store snapshot + rendered docs are removed too.
        (Any already-published Analysis graph is a separate copy and is untouched.)"""
        version = db.get_version(version_id)
        if version is None:
            raise ValueError(f"version {version_id} not found")
        try:
            self.store.delete_prefix(f"{version.project_id}/{version.dataset_id}/v{version.version_no}")
        except Exception as exc:
            logger.warning("delete_version: artifact cleanup failed for %s: %s", version_id, exc)
        deleted = db.delete_version(version_id)
        return {"deleted": deleted, "version_id": version_id}

    # ==================================================================
    # Clone — atomic copy of a (frozen) version into a new editable staging one
    # ==================================================================

    def clone_version(self, version_id: str) -> dict:
        """
        Deep-copy a version into a NEW staging version with fresh record IDs.

        Used to edit a frozen (main) version: the original stays an immutable
        snapshot; edits happen on the clone, which can then be re-promoted.
        Records, relationships, documents, and their validation/quality reports
        are copied and re-linked; artifacts are duplicated in the store.
        """
        src = db.get_version(version_id)
        if src is None:
            raise ValueError(f"version {version_id} not found")
        project = db.get_project(src.project_id)
        thr = project.min_threshold if project else 5

        new = db.create_version(src.project_id, src.dataset_id, note=f"clone of v{src.version_no}")

        # 1 · records (new IDs). Carry forward the ACCEPTED content
        # (staged + SME-approved) and RESET it to unverified, so the clone is a
        # fresh review cycle — SME/Quality/Validate all reflect an un-reviewed set.
        # Rejected / duplicate / SME-rejected records are left behind (the original
        # frozen version still retains them).
        records = [
            r for r in db.list_records(version_id)
            if r.status in (RecordStatus.STAGED, RecordStatus.SME_APPROVED)
        ]
        id_map: Dict[str, str] = {}
        new_records: List[SyntheticRecord] = []
        for r in records:
            nid = _new_id(r.element_type)
            id_map[r.id] = nid
            new_records.append(replace(
                r, id=nid, version_id=new.id, status=RecordStatus.STAGED, embedding_id=None,
                provenance={**r.provenance, "cloned_from": r.id},
            ))
        db.insert_records(new_records)

        # 2 · reports (remapped record IDs)
        reports = db.get_reports_for_version(version_id)
        val_reports: List[ValidationReport] = []
        qual_reports: List[QualityReport] = []
        for old_id, rep in reports.items():
            nid = id_map.get(old_id)
            if not nid:
                continue
            v = rep.get("validation")
            if v:
                val_reports.append(ValidationReport(
                    record_id=nid, schema_ok=v["schema_ok"], label_ok=v["label_ok"],
                    rules_ok=v["rules_ok"], reasons=v.get("reasons") or [],
                ))
            q = rep.get("quality")
            if q:
                qual_reports.append(QualityReport(
                    record_id=nid, realism=q["realism"], is_duplicate=q["is_duplicate"],
                    duplicate_of=id_map.get(q.get("duplicate_of")) if q.get("duplicate_of") else None,
                    near_dup_score=q.get("near_dup_score", 0.0), realism_notes=q.get("realism_notes", ""),
                ))
        db.upsert_validation_reports(val_reports)
        db.upsert_quality_reports(qual_reports)

        # 3 · relationships (remapped endpoints)
        new_rels: List[SyntheticRelationship] = []
        for rel in db.list_relationships(version_id):
            s, t = id_map.get(rel.source_record_id), id_map.get(rel.target_record_id)
            if not s or not t:
                continue
            new_rels.append(replace(
                rel, id=f"SREL_{uuid.uuid4().hex[:8].upper()}", version_id=new.id,
                source_record_id=s, target_record_id=t,
            ))
        db.insert_relationships(new_rels)

        # 4 · documents (remapped members + duplicated artifact)
        for d in db.list_documents(version_id):
            members = [id_map[x] for x in d.member_record_ids if x in id_map]
            sections = [
                {"heading": s.get("heading", ""),
                 "record_ids": [id_map[x] for x in s.get("record_ids", []) if x in id_map]}
                for s in d.sections
            ]
            ndoc = SyntheticDocument(
                id=f"SDOC_{uuid.uuid4().hex[:8].upper()}", project_id=d.project_id,
                version_id=new.id, doc_type=d.doc_type, title=d.title,
                member_record_ids=members, sections=sections, status=RecordStatus.STAGED,
                provenance={**d.provenance, "cloned_from": d.id},
            )
            try:
                src_key = d.provenance.get("artifact_key")
                if src_key and self.store.exists(src_key):
                    data = self.store.get_bytes(src_key)
                    nkey = self._doc_key(d.project_id, new.dataset_id, new.version_no, ndoc.id)
                    ndoc.artifact_uri = self.store.put_bytes(nkey, data, "text/markdown")
                    ndoc.provenance["artifact_key"] = nkey
            except Exception as exc:
                logger.warning("clone: artifact copy failed for %s: %s", d.id, exc)
            db.insert_document(ndoc)

        # 5 · stats + snapshot + lineage
        staged = [r for r in new_records if r.status in (RecordStatus.STAGED, RecordStatus.SME_APPROVED)]
        distribution = self.qual.compute_distribution(staged, new_rels, thr)
        artifact_uri = self._snapshot(src.project_id, new.dataset_id, new.version_no, staged, new_rels)
        stats = {
            "cloned_from": version_id, "cloned_from_version_no": src.version_no,
            "staged": len(staged), "relationships": len(new_rels),
            "documents": len(db.list_documents(new.id)), "distribution": distribution,
        }
        db.update_version_stats(new.id, stats, artifact_uri)
        db.add_lineage_edges(src.project_id, [(f"version:{version_id}", f"version:{new.id}", "cloned")])
        return {"version_id": new.id, "version_no": new.version_no, "cloned_from": version_id}

    def publish_to_analysis(
        self, version_id: str, workspace_id: str, progress_cb: ProgressCB = _noop,
    ) -> dict:
        """Convert accepted synthetic records into graph elements and ingest them
        into an Analysis workspace via the existing graph pipeline."""
        version = db.get_version(version_id)
        if version is None:
            raise ValueError(f"version {version_id} not found")

        # Publish everything accepted: SME-approved plus staged-but-unreviewed.
        # SME-rejected and duplicate records are excluded by virtue of their status.
        records = (
            db.list_records(version_id, status=RecordStatus.SME_APPROVED)
            + db.list_records(version_id, status=RecordStatus.STAGED)
        )
        if not records:
            raise ValueError("no accepted records to publish")

        progress_cb({"stage": "publish", "message": f"Publishing {len(records)} records to workspace"})

        # Group records into one synthetic ParsedDocument per doc_type.
        by_doc: Dict[DocumentType, List[SyntheticRecord]] = {}
        for r in records:
            by_doc.setdefault(r.doc_type, []).append(r)

        documents: List[ParsedDocument] = []
        elements: List[AtomicElement] = []
        rec_ids = {r.id for r in records}
        for dtype, recs in by_doc.items():
            doc_id = f"SYNDOC_{version_id[:8]}_{dtype.value}".replace(" ", "")
            body = "\n\n".join(r.text for r in recs)
            documents.append(ParsedDocument(
                id=doc_id, name=f"Synthetic {dtype.value} (v{version.version_no})",
                type=dtype, pages=[body], total_pages=1,
            ))
            for r in recs:
                el = AtomicElement(
                    id=r.id, type=r.element_type, text=r.text,
                    source=f"Synthetic {dtype.value} — {r.label}",
                    document_id=doc_id, confidence=1.0,
                    metadata={"section": r.label, "page_number": 1, "synthetic": True},
                )
                elements.append(el)

        # Relationships (positive only — negatives represent absence of an edge).
        relationships: List[Relationship] = []
        for rel in db.list_relationships(version_id):
            if not rel.is_positive:
                continue
            if rel.source_record_id in rec_ids and rel.target_record_id in rec_ids:
                relationships.append(Relationship(
                    source_id=rel.source_record_id, target_id=rel.target_record_id,
                    type=rel.rel_type, confidence=1.0, evidence=rel.rationale,
                ))

        # Lazy import to avoid an import cycle (deps → services → ...).
        from api.deps import get_graph_service
        gs = get_graph_service(workspace_id)
        doc_hashes = {d.id: f"synthetic-{version_id}-{d.id}" for d in documents}
        gs.build_knowledge_graph(documents, elements, relationships, doc_hashes)

        # Publication is NON-destructive: record statuses are left intact so the
        # version stays a reproducible snapshot and can be re-published (to the
        # same or another workspace). Publications are tracked on the version.
        summary = {
            "version_id": version_id, "workspace_id": workspace_id,
            "documents": len(documents), "elements": len(elements),
            "relationships": len(relationships),
            "nodes": gs.get_node_count(), "edges": gs.get_edge_count(),
        }
        stats = version.stats or {}
        pubs = list(stats.get("published_to", []))
        pubs.append({
            "workspace_id": workspace_id,
            "elements": len(elements), "relationships": len(relationships),
            "at": datetime.now(timezone.utc).isoformat(),
        })
        stats["published_to"] = pubs
        db.update_version_stats(version_id, stats)
        db.add_lineage_edges(version.project_id, [
            (f"version:{version_id}", f"workspace:{workspace_id}", "published"),
        ])

        progress_cb({"stage": "complete", "message": "Published to Analysis", "summary": summary})
        return summary

    def publish_documents_to_store(self, version_id: str, progress_cb: ProgressCB = _noop) -> dict:
        """Push SME-approved (+ staged-but-unreviewed) whole documents into the
        shared, cross-workspace document store — tagged `_gen`. Does NOT touch
        any Analysis workspace directly; a workspace pulls from the store on
        demand via a separate import step (see backend/services/synthetic_import.py)."""
        version = db.get_version(version_id)
        if version is None:
            raise ValueError(f"version {version_id} not found")

        docs = [
            d for d in db.list_documents(version_id)
            if d.status in (RecordStatus.SME_APPROVED, RecordStatus.STAGED)
        ]
        if not docs:
            raise ValueError("no accepted documents to publish")

        progress_cb({"stage": "publish", "message": f"Publishing {len(docs)} documents to the store"})

        published = []
        for doc in docs:
            key = doc.provenance.get("artifact_key") or self._doc_key(
                doc.project_id, version.dataset_id, version.version_no, doc.id,
            )
            if not self.store.exists(key):
                # Shouldn't normally happen — generation always writes the artifact —
                # but fall back to re-rendering so publish never silently drops content.
                key = self._doc_key(doc.project_id, version.dataset_id, version.version_no, doc.id)
                self.store.put_text(key, self.document_markdown(version_id, doc.id), "text/markdown")
            entry = db.publish_document_to_store(doc.project_id, version_id, doc, key)
            db.update_document_content(doc.id, status=RecordStatus.PUBLISHED)
            published.append(entry.to_dict())

        stats = version.stats or {}
        pubs = list(stats.get("published_to_store", []))
        pubs.append({"count": len(published), "at": datetime.now(timezone.utc).isoformat()})
        stats["published_to_store"] = pubs
        db.update_version_stats(version_id, stats)
        db.add_lineage_edges(version.project_id, [
            (f"version:{version_id}", "store:documents", "published_to_store"),
        ])

        summary = {"version_id": version_id, "published": len(published), "documents": published}
        progress_cb({"stage": "complete", "message": "Published to document store", "summary": summary})
        return summary

    def publish_documents(self, document_ids: List[str]) -> dict:
        """Publish an explicit set of documents by id — the professional-UI
        Document Library's "Send to Document Storage" action. Version-agnostic:
        each document's version is resolved internally purely to reconstruct
        its artifact key, never surfaced to the caller."""
        if not document_ids:
            raise ValueError("no documents selected")

        published = []
        by_project: Dict[str, List[str]] = {}
        for doc_id in document_ids:
            doc = db.get_document(doc_id)
            if doc is None:
                raise ValueError(f"document {doc_id} not found")
            if doc.status != RecordStatus.SME_APPROVED:
                raise ValueError(f"document {doc_id} must be approved before publishing")
            version = db.get_version(doc.version_id) if doc.version_id else None
            if version is None:
                raise ValueError(f"document {doc_id} has no associated version")

            key = doc.provenance.get("artifact_key") or self._doc_key(
                doc.project_id, version.dataset_id, version.version_no, doc.id,
            )
            if not self.store.exists(key):
                key = self._doc_key(doc.project_id, version.dataset_id, version.version_no, doc.id)
                self.store.put_text(key, self.document_markdown(doc.version_id, doc.id), "text/markdown")
            entry = db.publish_document_to_store(doc.project_id, doc.version_id, doc, key)
            db.update_document_content(doc.id, status=RecordStatus.PUBLISHED)
            published.append(entry.to_dict())
            by_project.setdefault(doc.project_id, []).append(doc.id)

        for project_id, ids in by_project.items():
            db.add_lineage_edges(project_id, [
                (f"document:{d}", "store:documents", "published_to_store") for d in ids
            ])

        return {"published": len(published), "documents": published}

    def recall_document(self, document_id: str) -> dict:
        """Bring a published document back from the store for further editing.

        Removes it from the shared document store and reverts its status to
        SME_APPROVED (editable + re-publishable). Copies already imported into
        Analysis workspaces are independent and are left untouched — the returned
        ``imported_into`` lets the UI warn about them."""
        doc = db.get_document(document_id)
        if doc is None:
            raise ValueError(f"document {document_id} not found")
        if doc.status != RecordStatus.PUBLISHED:
            raise ValueError(f"document {document_id} is not published")

        result = db.recall_document_from_store(document_id)
        db.add_lineage_edges(doc.project_id, [
            (f"document:{document_id}", "store:documents", "recalled_from_store"),
        ])
        db.touch_project(doc.project_id)
        return {
            "document_id": document_id,
            "status": RecordStatus.SME_APPROVED.value,
            "removed_from_store": result["removed"],
            "imported_into": result["imported_into"],
        }

    # ==================================================================
    # Lineage + artifacts
    # ==================================================================

    def lineage(self, project_id: str) -> dict:
        return {"project_id": project_id, "edges": db.list_lineage(project_id)}

    # ==================================================================
    # Export — datasets + draft documents
    # ==================================================================

    def export_records_jsonl(self, version_id: str) -> str:
        """Authoritative records export, regenerated from the DB (reflects edits)."""
        recs = db.list_records(version_id)
        return "\n".join(json.dumps(r.to_dict(), ensure_ascii=False) for r in recs)

    def export_relationships_jsonl(self, version_id: str) -> str:
        rels = db.list_relationships(version_id)
        return "\n".join(json.dumps(r.to_dict(), ensure_ascii=False) for r in rels)

    def _find_document(self, version_id: str, doc_id: str) -> SyntheticDocument:
        for d in db.list_documents(version_id):
            if d.id == doc_id:
                return d
        raise ValueError(f"document {doc_id} not found in version {version_id}")

    def document_markdown(self, version_id: str, doc_id: str) -> str:
        """Return the draft document's Markdown — from the stored artifact, or
        regenerated from its sections if the artifact is unavailable."""
        doc = self._find_document(version_id, doc_id)
        key = doc.provenance.get("artifact_key")
        if key:
            try:
                if self.store.exists(key):
                    return self.store.get_bytes(key).decode("utf-8")
            except Exception as exc:
                logger.warning("export md: artifact read failed (%s) — regenerating", exc)
        # Fallback: rebuild from sections — direct body text (document-level
        # generation) or, for the legacy element-assembled path, record texts.
        texts = {r.id: r.text for r in db.list_records(version_id)}
        lines = [f"# {doc.title}", ""]
        for s in doc.sections:
            lines.append(f"## {s.get('heading', 'Section')}")
            if s.get("body"):
                lines.append(s["body"])
            else:
                for i, rid in enumerate(s.get("record_ids", []), 1):
                    if rid in texts:
                        lines.append(f"{i}. {texts[rid]}")
            lines.append("")
        return "\n".join(lines)

    def document_docx(self, version_id: str, doc_id: str) -> bytes:
        """Render the draft document as a Word .docx."""
        from docx import Document as Docx  # python-docx (already a dependency)
        doc = self._find_document(version_id, doc_id)
        texts = {r.id: r.text for r in db.list_records(version_id)}
        d = Docx()
        d.add_heading(doc.title, level=0)
        if doc.provenance.get("brief"):
            d.add_paragraph("Generated to a user brief.").italic = True
        for s in doc.sections:
            d.add_heading(s.get("heading", "Section"), level=1)
            if s.get("body"):
                d.add_paragraph(s["body"])
            else:
                for rid in s.get("record_ids", []):
                    if rid in texts:
                        d.add_paragraph(texts[rid], style="List Number")
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()

    def export_documents_zip(
        self, version_id: str, doc_ids: Optional[List[str]] = None, fmt: str = "md",
    ) -> bytes:
        """A lean ZIP of just the documents themselves (.md or .docx) — no
        records.jsonl/relationships.jsonl/manifest.json. ``doc_ids`` narrows
        to a specific selection; omit it to include every document in the
        version."""
        all_docs = db.list_documents(version_id)
        docs = [d for d in all_docs if d.id in set(doc_ids)] if doc_ids else all_docs
        if not docs:
            raise ValueError("no documents to export")

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            seen_names: Dict[str, int] = {}
            for d in docs:
                safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in d.title)[:60] or d.id
                n = seen_names.get(safe, 0)
                seen_names[safe] = n + 1
                name = safe if n == 0 else f"{safe}_{n}"
                if fmt == "docx":
                    z.writestr(f"{name}.docx", self.document_docx(version_id, d.id))
                else:
                    z.writestr(f"{name}.md", self.document_markdown(version_id, d.id))
        return buf.getvalue()

    def export_bundle_zip(self, version_id: str) -> bytes:
        """A single ZIP: records + relationships JSONL, every draft doc (.md),
        and a manifest.json (version stats + project + lineage)."""
        version = db.get_version(version_id)
        if version is None:
            raise ValueError(f"version {version_id} not found")
        project = db.get_project(version.project_id)
        docs = db.list_documents(version_id)

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr("records.jsonl", self.export_records_jsonl(version_id))
            z.writestr("relationships.jsonl", self.export_relationships_jsonl(version_id))
            for d in docs:
                safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in d.title)[:60] or d.id
                z.writestr(f"docs/{safe}.md", self.document_markdown(version_id, d.id))
            manifest = {
                "version": version.to_dict(),
                "project": project.to_dict() if project else None,
                "record_count": len(db.list_records(version_id)),
                "document_count": len(docs),
                "lineage": db.list_lineage(version.project_id),
            }
            z.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
        return buf.getvalue()

    def _snapshot(
        self, project_id: str, dataset_id: str, version_no: int,
        records: List[SyntheticRecord], relationships: List[SyntheticRelationship],
    ) -> str:
        base = f"{project_id}/{dataset_id}/v{version_no}"
        rec_jsonl = "\n".join(json.dumps(r.to_dict(), ensure_ascii=False) for r in records)
        uri = self.store.put_text(f"{base}/records.jsonl", rec_jsonl, "application/x-ndjson")
        rel_jsonl = "\n".join(json.dumps(r.to_dict(), ensure_ascii=False) for r in relationships)
        self.store.put_text(f"{base}/relationships.jsonl", rel_jsonl, "application/x-ndjson")
        return uri

    def _doc_key(self, project_id: str, dataset_id: str, version_no: int, doc_id: str) -> str:
        return f"{project_id}/{dataset_id}/v{version_no}/docs/{doc_id}.md"

    def _lineage_for_generation(
        self, project_id: str, version_id: str, selections: List[dict],
        staged: List[SyntheticRecord],
    ) -> None:
        edges = [(f"project:{project_id}", f"version:{version_id}", "generated")]
        for sel in selections:
            target = sel.get("cell") or sel.get("element_type") or "records"
            edges.append((f"version:{version_id}", f"target:{target}", "targets"))
        db.add_lineage_edges(project_id, edges)
