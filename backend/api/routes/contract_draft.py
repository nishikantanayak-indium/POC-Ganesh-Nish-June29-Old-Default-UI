"""Contract Draft generation — workspace-scoped, SSE streaming.

Mirrors pipeline.py's _stream_pipeline event shape (step_start/step_complete/
pipeline_complete/error via the shared _sse() convention) so the frontend's
already-proven SSE handling code applies here with a new step vocabulary."""
from __future__ import annotations

import asyncio
import json
import time
from typing import AsyncGenerator, Optional

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel

from api.deps import get_contract_draft_service, get_graph_service
from db import postgres as db
from services.contract_draft_service import _quote_appears_in

router = APIRouter(prefix="/api/workspaces/{workspace_id}")


def _sse(data: dict) -> str:
    return f"data: {json.dumps(data)}\n\n"


class GenerateDraftBody(BaseModel):
    template: str = "rfp_mirror"  # 'rfp_mirror' | 'services_agreement' | 'rfp_response'


async def _stream_generate_draft(workspace_id: str, template: str) -> AsyncGenerator[str, None]:
    t0 = time.perf_counter()
    svc = get_contract_draft_service()
    gs = get_graph_service(workspace_id)

    yield _sse({"stage": "queued", "message": "Draft generation started"})
    await asyncio.sleep(0)

    yield _sse({"stage": "grounding", "message": "Assessing coverage and collecting requirement/risk evidence…"})
    await asyncio.sleep(0)
    try:
        bundle = await asyncio.to_thread(svc.build_grounding_bundle, gs)
    except Exception as exc:
        yield _sse({"stage": "error", "message": str(exc)})
        return
    yield _sse({
        "stage": "grounding",
        "message": (
            f"{bundle['summary']['requirements_needing_attention']} requirement(s) to address, "
            f"{bundle['summary']['gaps_count']} gap(s) flagged"
        ),
    })
    await asyncio.sleep(0)

    yield _sse({"stage": "drafting", "message": f"Drafting document ({template})…"})
    await asyncio.sleep(0)
    try:
        title, sections, unresolved_gaps = await asyncio.to_thread(svc.generate_draft, gs, template, bundle)
    except Exception as exc:
        yield _sse({"stage": "error", "message": str(exc)})
        return
    yield _sse({"stage": "drafting", "message": f"Drafted '{title}' — {len(sections)} section(s)"})
    await asyncio.sleep(0)

    total_citations = sum(len(s["citations"]) for s in sections)
    strong = sum(1 for s in sections for c in s["citations"] if c["verdict"] == "strong")
    yield _sse({
        "stage": "citing",
        "message": f"Verified {total_citations} citation(s) against drafted text — {strong} strong",
    })
    await asyncio.sleep(0)

    yield _sse({"stage": "persisting", "message": "Saving draft…"})
    await asyncio.sleep(0)
    draft = await asyncio.to_thread(
        db.create_contract_draft, workspace_id, title, template, sections, unresolved_gaps, bundle["summary"],
    )

    yield _sse({
        "stage": "done",
        "message": "Draft ready for review",
        "summary": {
            "draft_id": draft.id, "title": draft.title, "sections": len(sections),
            "gaps": len(unresolved_gaps), "elapsed": round(time.perf_counter() - t0, 2),
        },
    })


@router.post("/draft/generate")
async def generate_draft(workspace_id: str, body: GenerateDraftBody) -> StreamingResponse:
    return StreamingResponse(
        _stream_generate_draft(workspace_id, body.template),
        media_type="text/event-stream",
    )


@router.get("/drafts")
async def list_drafts(workspace_id: str) -> dict:
    drafts = await asyncio.to_thread(db.list_contract_drafts, workspace_id)
    return {"drafts": [d.to_dict() for d in drafts]}


_TEMPLATE_CHOICES = [
    ("rfp_mirror", "Mirror the RFP"),
    ("services_agreement", "Services Agreement"),
    ("rfp_response", "RFP Response"),
]


@router.get("/draft/templates")
async def get_draft_templates(workspace_id: str) -> dict:
    """No-LLM structure preview for the template picker — shows the real
    section list a template would produce (for 'rfp_mirror', headings
    detected out of the actually-ingested RFP), not just a text blurb."""
    svc = get_contract_draft_service()
    gs = get_graph_service(workspace_id)

    def _build() -> list[dict]:
        out = []
        for value, label in _TEMPLATE_CHOICES:
            sections = svc.preview_template_sections(gs, value)
            out.append({"value": value, "label": label, "sections": sections})
        return out

    return {"templates": await asyncio.to_thread(_build)}


@router.get("/draft/{draft_id}")
async def get_draft(workspace_id: str, draft_id: str) -> dict:
    draft = await asyncio.to_thread(db.get_contract_draft, draft_id)
    if draft is None or draft.workspace_id != workspace_id:
        raise HTTPException(status_code=404, detail="Draft not found")
    return draft.to_dict()


class UpdateDraftBody(BaseModel):
    status: Optional[str] = None
    sections: Optional[list] = None


@router.patch("/draft/{draft_id}")
async def update_draft(workspace_id: str, draft_id: str, body: UpdateDraftBody) -> dict:
    existing = await asyncio.to_thread(db.get_contract_draft, draft_id)
    if existing is None or existing.workspace_id != workspace_id:
        raise HTTPException(status_code=404, detail="Draft not found")

    sections = body.sections
    if sections is not None:
        # A manual or AI-assisted edit can change the body enough that a
        # previously-verified citation quote no longer actually appears in
        # it — re-check now rather than let a stale "strong" citation imply
        # evidence that no longer exists in the edited text.
        for s in sections:
            for c in s.get("citations", []):
                if c.get("quote") and not _quote_appears_in(c["quote"], s.get("body", "")):
                    c["quote"], c["verdict"] = "", "weak"

    updated = await asyncio.to_thread(
        db.update_contract_draft, draft_id, status=body.status, sections=sections,
    )
    return updated.to_dict()


class ReviseSectionBody(BaseModel):
    selected_text: str
    instruction: str


@router.post("/draft/{draft_id}/sections/{section_index}/revise")
async def revise_section(workspace_id: str, draft_id: str, section_index: int, body: ReviseSectionBody) -> dict:
    """AI-assisted revision of a specific selected span within one section —
    distinct from a full manual rewrite. Returns only the replacement text;
    the client splices it into the section body at the selection offsets and
    saves via the regular PATCH, so citation re-verification above applies."""
    draft = await asyncio.to_thread(db.get_contract_draft, draft_id)
    if draft is None or draft.workspace_id != workspace_id:
        raise HTTPException(status_code=404, detail="Draft not found")
    if section_index < 0 or section_index >= len(draft.sections):
        raise HTTPException(status_code=404, detail="Section not found")
    if not body.selected_text.strip() or not body.instruction.strip():
        raise HTTPException(status_code=400, detail="selected_text and instruction are required")

    svc = get_contract_draft_service()
    revised = await asyncio.to_thread(
        svc.revise_selection, draft.sections[section_index]["body"], body.selected_text, body.instruction,
    )
    return {"revised_text": revised}


@router.get("/draft/{draft_id}/export.md")
async def export_draft_markdown(workspace_id: str, draft_id: str) -> Response:
    draft = await asyncio.to_thread(db.get_contract_draft, draft_id)
    if draft is None or draft.workspace_id != workspace_id:
        raise HTTPException(status_code=404, detail="Draft not found")
    lines = [f"# {draft.title}", ""]
    for s in draft.sections:
        lines.append(f"## {s['heading']}")
        lines.append(s["body"])
        lines.append("")
    markdown = "\n".join(lines)
    return Response(
        content=markdown, media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{draft.title[:60]}.md"'},
    )


@router.get("/draft/{draft_id}/export.docx")
async def export_draft_docx(workspace_id: str, draft_id: str) -> Response:
    draft = await asyncio.to_thread(db.get_contract_draft, draft_id)
    if draft is None or draft.workspace_id != workspace_id:
        raise HTTPException(status_code=404, detail="Draft not found")

    def _build() -> bytes:
        import io
        from docx import Document as Docx

        doc = Docx()
        doc.add_heading(draft.title, level=0)
        for s in draft.sections:
            doc.add_heading(s["heading"], level=1)
            for para in s["body"].split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    content = await asyncio.to_thread(_build)
    return Response(
        content=content, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{draft.title[:60]}.docx"'},
    )
